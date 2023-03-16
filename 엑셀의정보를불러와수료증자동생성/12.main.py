# 엑셀의 정보를 불러와 수료증 자동 생성

## pip install openpyxl - 엑셀을 사용하기 위한 라이브러리
## pip install python-docx - 워드를 사용하기 위한 라이브러리
## pip install docx2pdf - 워드를 PDF로 변환할 때 사용사는 라이브러리


import pandas as pd



### 수료증 명단 엑셀 파일 만들기
df = pd.DataFrame([
    ['홍길동', '1990.01.02', '2023-0001'],
    ['김철수', '2000.08.08', '2023-0002'],
    ['김영희', '2000.09.09', '2023-0003'],
    ['이서준', '2010.10.10', '2023-0004'],
    ['장다인', '2017.12.12', '2023-0005'],
    ['김민준', '1990.05.06', '2023-0006'],
    ])

print(df)

#df.to_excel(r'04_AutomationProgram\pjt12_엑셀의정보를불러와수료증자동생성\result\수료증명단.xlsx', index=False, header=False)


### 판다스 라이브러리로 값을 엑셀로 저장 후 불러오는 코드 만들기

from openpyxl import load_workbook  # 엑셀에서 파일을 읽기위해서 openpyxl 라이브러리에서 load_workbook을 불러온다

load_wb = load_workbook(r'04_AutomationProgram\pjt12_엑셀의정보를불러와수료증자동생성\result\수료증명단.xlsx')
load_ws = load_wb.active

name_list =[]
birthday_list = []
ho_list = []
for i in range(1, load_ws.max_row +1):
    name_list.append(load_ws.cell(i,1).value)
    birthday_list.append(load_ws.cell(i,2).value)
    ho_list.append(load_ws.cell(i,3).value)

print(name_list)
print(birthday_list)
print(ho_list)


###수료증 레이아웃을 워드 문서로 만들기
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def GenerateCertificate():
    
    doc = docx.Document()

    style = doc.styles['Normal']
    style.font.name = '맑은고딕'
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\r\n')
    run = para.add_run('제 2020-00001 호\n')
    run.font.name = '맑은고딕'
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    run.bold = True
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 성 명 : 장다인\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 생 년 월 일 : 2017.12.12\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 과 정 : 파이썬과 40개의 작품들\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 날 짜 : 2021.08.05~2021.09.09\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 위 사람은 파이썬과 40개의 작품들 교육과정을\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 이수하였으므로 이 증서를 수여 합니다.\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('2022.08.19')
    run.font.size= docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('파이썬교육기관장')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.save(r'04_AutomationProgram\pjt12_엑셀의정보를불러와수료증자동생성\result\수료증결과.docx')
   
GenerateCertificate()


def GenerateCertificate(name, birthday, num):
    doc = docx.Document()
    
    style = doc.styles['Normal']
    style.font.name = '맑은고딕'
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\r\n')
    run = para.add_run('제 ' + num +' 호\n')
    run.font.name = '맑은고딕'
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    run.bold = True
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 성 명 : '+name+'\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 생 년 월 일 : '+ birthday +'\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 과 정 : 파이썬과 40개의 작품들\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 날 짜 : 2021.08.05~2021.09.09\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 위 사람은 파이썬과 40개의 작품들 교육과정을\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 이수하였으므로 이 증서를 수여 합니다.\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('2022.08.19')
    run.font.size= docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('파이썬교육기관장')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc
    
    


### 수료증 생성 후 PDF로 변환하는 코드 만들기
from docx2pdf import convert

for i in range(len(name_list)):
    doc = GenerateCertificate(name_list[i], birthday_list[i], ho_list[i]);
    savePath = r'04_AutomationProgram\pjt12_엑셀의정보를불러와수료증자동생성\result\수료증결과_'+ name_list[i]
    doc.save(savePath+'.docx')
    convert(savePath+'.docx', savePath+'.pdf')